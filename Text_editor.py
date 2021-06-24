from tkinter import *
from tkinter.filedialog import askopenfilename, asksaveasfilename
from tkinter.messagebox import showerror
from Dialog_window import *
#from Image_dialog import *
import cv2
import pytesseract
import docx
import os

class Editor(Frame):
    def __init__(self, root):
        self.root = root
        self.root.title("Prozor za obradu teksta")
        self.root.iconbitmap(r'C:\Users\Korisnik\Desktop\Zavrsni_rad\icon.ico')
        super(Editor, self).__init__(self.root)
        self.KreirajSučelje()
        self.pack(fill=BOTH, expand=True)
        return

    def KreirajSučelje(self):
        #ovo je izgled editora i pokretna traka na rubovima
        self.Tn = 1
        ys = Scrollbar(self.root, orient=VERTICAL)
        ys.pack(side=RIGHT, fill=Y)
        xs = Scrollbar(self.root, orient=HORIZONTAL)
        xs.pack(side=BOTTOM, fill=X)

        self.T = Text(self, wrap = NONE, spacing1 = 5)
        self.T.pack(fill=BOTH, expand=True)
        self.T.config(xscrollcommand = xs.set, yscrollcommand = ys.set)
        self.T.focus_set()
        self.F = None

        ys.config(command=self.T.yview)
        xs.config(command=self.T.xview)

        #ovo je izbornik
        mB = Menu(self.root)
        mD = Menu(mB, tearoff = 0)
        mD.add_command(label="Učitaj fotografiju", underline=0, accelerator="Ctrl+U", command=self.Učitaj_sliku)
        self.root.bind("<Control-u>", self.Učitaj_sliku)
        mD.add_separator()
        mD.add_command(label="Nova", underline=0, accelerator="Ctrl+N", command=self.Nova)
        self.root.bind("<Control-n>", self.Nova)
        mD.add_command(label="Otvori", underline=0, accelerator="Ctrl+O", command=self.Otvori)
        self.root.bind("<Control-o>", self.Otvori)
        mD.add_command(label="Spremi", underline=0, accelerator="Ctrl+S", command=self.Spremi)
        self.root.bind("<Control-s>", self.Spremi)
        mD.add_command(label="Spremi kao...", underline=7, accelerator="Ctrl+Shift+S", command=self.SpremiKao)
        self.root.bind("<Control-Shift-s>", self.SpremiKao)
        mD.add_separator()
        mD.add_command(label="Kraj", underline=0, accelerator="Ctrl+Q", command=self.Kraj)
        self.root.bind("<Control-q>", self.Kraj)
        mB.add_cascade(menu=mD, label="Datoteka")

        uM = Menu(mD, tearoff=0)
        uM.add_command(label="Font", underline=0, accelerator="Ctrl+F", command=self.Font)
        self.root.bind("<Control-f>", self.Font)
        uM.add_separator()
        uM.add_command(label="Izreži", underline=0, accelerator="Ctrl+X", command=self.Izreži)
        #self.root.bind("<Control-x>", self.Izreži)
        uM.add_command(label="Kopiraj", underline=0, accelerator="Ctrl+C", command=self.Kopiraj)
        #self.root.bind("<Control-c>", self.Kopiraj)
        uM.add_command(label="Zalijepi", underline=0, accelerator="Ctrl+V", command=self.Zalijepi)
        #self.root.bind("<Control-v>", self.Zalijepi)
        mB.add_cascade(menu=uM, label="Uređivanje")

        self.root.config(menu = mB)
        return


    def Kraj(self, e = None):
        self.root.destroy()
        return

    def Nova(self, e = None):
        self.T.delete(1.0, END)
        self.F = None
        self.root.title("Bez imena")
        return

    def Otvori(self, e = None):
        fname = askopenfilename(filetype=[("Sve datoteke", "*.*"),("Text datoteke", "*.txt"), ("Docx datoteke", "*.docx")], title="Odaberi datoteku")
        if fname[-5:] == ".docx" or fname[-4:] == ".doc":
            try:
                text = []
                doc = docx.Document(fname)
                all = doc.paragraphs
                for red in all:
                    text.append(red.text)
                    data = "\n".join(text)
                self.T.insert(0.0, data)
                self.root.title(fname)
                self.F = fname
            except:
                showerror("Program za uređivanje teksta", "Datoteka ne postoji")
        elif fname:
            try:
                #sljedeca linija brise prethodno zapisani tekst, bez njega novi tekst se dodaje na prethodni
                #self.T.delete(1.0, END)
                with open(fname, "r") as fread:
                    p = fread.read()
                    self.T.insert(0.0, p)
                    self.root.title(fname)
                    self.F = fname
            except:
                showerror("Program za uređivanje teksta", "Datoteka ne postoji")
        return

    def Spremi(self, e = None):
        if self.F == None:
            self.SpremiKao()
        else:
            base = os.path.basename(self.F)
            name = os.path.splitext(base)[0]
            ext = os.path.splitext(base)[1]
            if ext == ".docx" or ext == ".doc":
                doc = docx.Document()
                doc.add_heading("Dokument - " + str(name), 0)
                p = doc.add_paragraph(str(self.T.get(0.0, END)))
                doc.save(self.F)
            else:
                with open(self.F, "w") as f:
                    f.write(self.T.get(0.0, END))
        return

    def SpremiKao(self, e = None):
        #fname = asksaveasfilename(filetype=[("Datoteka python", "*.py"), ("Tekstualne datoteke", "*.txt"),
        #                                  ("Sve datoteke", "*.*"), ("Png slika", "*.png"),
        #                                  ("jpg slika", "*.jpg")], defaultextension=".txt", title="Spremi kao...")

        fname = asksaveasfilename(filetype=[("Sve datoteke", "*.*"), ("Text datoteke", "*.txt"), ("Docx datoteke", "*.docx")], defaultextension=".txt", title="Spremi kao...")

        if fname[-5:] == ".docx" or fname[-4:] == ".doc":
            self.F = fname
            self.Spremi()
            self.root.title(self.F)
        elif fname:
            self.F = fname
            self.Spremi()
            self.root.title(fname)
        return

    def Kopiraj(self, e = None):
        try:
            self.root.clipboard_clear()
            self.root.clipboard_append(self.T.get(SEL_FIRST, SEL_LAST))
        except:
            showerror("Prgram za uređivanje teksta", "Niste označili tekst")
        return

    def Izreži(self, e = None):
        try:
            self.Kopiraj()
            self.T.delete(SEL_FIRST, SEL_LAST)
        except:
            return
        return

    def Zalijepi(self, e = None):
        try:
            self.T.insert(INSERT, self.root.clipboard_get())
        except:
            showerror("Prgram za uređivanje teksta", "Međuspremnik je prazan")
        return

    def Font(self, e = None):
        try:
            f = FontDialog(self.root)
            if f.Font:
                t = "t_" + str(self.Tn)
                self.T.tag_add(t, SEL_FIRST, SEL_LAST)
                self.T.tag_config(t, font= f.Font[0], foreground=f.Font[1])
                self.Tn += 1
        except:
            return
        return

    def mouse_crop(self, event, x, y, flags, param):

        # kada se klikne lijevi klik misa uzimaju se x, y koordinate za sjeckanje slike
        if event == cv2.EVENT_LBUTTONDOWN:
            self.x_start, self.y_start, self.x_end, self.y_end = x, y, x, y
            self.cropping = True

        # prati pokrete misa kada se drzi prst na lijevom kliku
        elif event == cv2.EVENT_MOUSEMOVE:
            if self.cropping == True:
                self.x_end, self.y_end = x, y

        # kada se pusti lijevi klik zaustavlja sjeckanje
        elif event == cv2.EVENT_LBUTTONUP:
            self.x_end, self.y_end = x, y
            self.cropping = False  # cropping is finished
            refPoint = [(self.x_start, self.y_start), (self.x_end, self.y_end)]

            if len(refPoint) == 2:
                roi = self.oriImage[refPoint[0][1]:refPoint[1][1], refPoint[0][0]:refPoint[1][0]]
                pytesseract.pytesseract.tesseract_cmd = "C:\\Program Files\\Tesseract-OCR\\tesseract.exe"
                img = cv2.cvtColor(roi, cv2.COLOR_BGR2RGB)
                custom_config = r'-l hrv --psm 6'
                rez = pytesseract.image_to_string(img, lang="hrv", config=custom_config)
                self.T.insert(0.0, rez)

                # <---------------nije potrebno dodatno crtati kvadrate kada u while to radimo----------->
                # boxes = pytesseract.image_to_data(img)
                # for x, b in enumerate(boxes.splitlines()):
                #    if x != 0:
                #        b = b.split()
                #        if len(b) == 12:
                #            x, y, w, h = int(b[6]), int(b[7]), int(b[8]), int(b[9])
                #            cv2.rectangle(roi, (x, y), (w + x, h + y), (0, 0, 255), 1)
                #            cv2.putText(img, b[11], (x, y + 65), cv2.FONT_HERSHEY_COMPLEX, 1, (255, 0, 0), 2)
                # cv2.imshow("Odabrani tekst za čitanje", roi)
                return

    def Učitaj_sliku(self, e = None):
        #ImageDialog je novi dialog prikaza slike....(ne koristi se pošto openCV to radi)
        #try:
        #    ImageDialog(self.root)
        #except:
        #    print("ne radi")
        #return
        fname = askopenfilename(
            filetype=[("Fotografija datoteke", "*.jpg;" "*.jpeg;" "*.png"), ("PDF datoteke", "*.pdf")],
            title="Odaberi fotografiju")

        self.cropping = False
        self.x_start, self.y_start, self.x_end, self.y_end = 0, 0, 0, 0
        self.image = cv2.imread(fname)
        self.oriImage = self.image.copy()

        cv2.namedWindow("Slika")
        cv2.setMouseCallback("Slika", self.mouse_crop)
        while True:
            i = self.image.copy()
            if cv2.waitKey(10) == 27:
                break
            elif not self.cropping:
                cv2.imshow("Slika", self.image)
            elif self.cropping:
                cv2.rectangle(i, (self.x_start, self.y_start), (self.x_end, self.y_end), (0, 255, 0), 2)
                cv2.imshow("Slika", i)
            return

        cv2.destroyAllWindows()

def main():
    e = Editor(Tk())
    mainloop()
main()